import streamlit as st
import logging
import tempfile
import os
from typing import Optional
import PyPDF2
import docx
import io

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_file(uploaded_file) -> Optional[str]:
    """Extract text from uploaded file (PDF, DOCX, TXT) with robust error handling"""
    
    if not uploaded_file:
        logger.error("No file provided for text extraction")
        st.error("❌ No file provided for text extraction")
        return None
    
    try:
        filename = uploaded_file.name.lower()
        logger.info(f"Extracting text from {filename}")
        
        # Get file content - FIXED: Proper error handling
        try:
            file_content = uploaded_file.getvalue()
            logger.info(f"Successfully read file content: {len(file_content)} bytes")
        except Exception as e:
            logger.error(f"Failed to read file content from {uploaded_file.name}: {str(e)}")
            st.error(f"❌ Could not read file {uploaded_file.name}: {str(e)}")
            return None
        
        if not file_content:
            logger.error(f"File {filename} is empty")
            st.error(f"❌ File {uploaded_file.name} appears to be empty")
            return None
        
        if filename.endswith('.pdf'):
            return extract_text_from_pdf(file_content, uploaded_file.name)
        elif filename.endswith('.docx'):
            return extract_text_from_docx(file_content, uploaded_file.name)
        elif filename.endswith('.txt'):
            return extract_text_from_txt(file_content, uploaded_file.name)
        else:
            logger.error(f"Unsupported file format: {filename}")
            st.error(f"❌ Unsupported file format: {filename}")
            return None
            
    except Exception as e:
        logger.error(f"Critical error extracting text from {uploaded_file.name}: {str(e)}")
        st.error(f"❌ Critical error reading file {uploaded_file.name}: {str(e)}")
        return None

def extract_text_from_pdf(file_content: bytes, filename: str) -> Optional[str]:
    """Extract text from PDF file with multiple extraction methods"""
    try:
        logger.info(f"Starting PDF text extraction for {filename}")
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        
        if len(pdf_reader.pages) == 0:
            logger.warning(f"PDF {filename} has no pages")
            st.error(f"❌ PDF {filename} has no pages")
            return None
        
        text_parts = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
                    logger.info(f"Extracted {len(page_text)} chars from page {page_num + 1}")
                else:
                    logger.warning(f"No text found on page {page_num + 1} of {filename}")
            except Exception as e:
                logger.error(f"Error extracting from page {page_num + 1} of {filename}: {str(e)}")
                continue
        
        if not text_parts:
            logger.error(f"No text could be extracted from PDF {filename}")
            st.error(f"❌ Could not extract text from PDF {filename}. The file may be password-protected or contain only images.")
            return None
        
        full_text = '\n\n'.join(text_parts)
        logger.info(f"Successfully extracted {len(full_text)} characters from PDF {filename}")
        
        # ADDED: Show success message to user
        st.success(f"✅ Successfully extracted text from PDF ({len(full_text)} characters)")
        
        return full_text
        
    except Exception as e:
        logger.error(f"PDF extraction failed for {filename}: {str(e)}")
        st.error(f"❌ Error reading PDF {filename}: {str(e)}")
        return None

def extract_text_from_docx(file_content: bytes, filename: str) -> Optional[str]:
    """Extract text from DOCX file"""
    try:
        logger.info(f"Starting DOCX text extraction for {filename}")
        doc = docx.Document(io.BytesIO(file_content))
        
        text_parts = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        if not text_parts:
            logger.error(f"No text found in DOCX {filename}")
            st.error(f"❌ No text content found in {filename}")
            return None
        
        full_text = '\n\n'.join(text_parts)
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX {filename}")
        
        # ADDED: Show success message to user
        st.success(f"✅ Successfully extracted text from DOCX ({len(full_text)} characters)")
        
        return full_text
        
    except Exception as e:
        logger.error(f"DOCX extraction failed for {filename}: {str(e)}")
        st.error(f"❌ Error reading DOCX {filename}: {str(e)}")
        return None

def extract_text_from_txt(file_content: bytes, filename: str) -> Optional[str]:
    """Extract text from TXT file with encoding detection"""
    
    logger.info(f"Starting TXT text extraction for {filename}")
    
    # Common encodings to try
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'ascii']
    
    for encoding in encodings:
        try:
            text = file_content.decode(encoding)
            if text and text.strip():
                logger.info(f"Successfully decoded TXT {filename} with {encoding} encoding")
                
                # ADDED: Show success message to user
                st.success(f"✅ Successfully extracted text from TXT ({len(text)} characters)")
                
                return text.strip()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.warning(f"Error with {encoding} encoding for {filename}: {str(e)}")
            continue
    
    logger.error(f"Could not decode TXT file {filename} with any encoding")
    st.error(f"❌ Could not read text file {filename}. Please check the file encoding.")
    return None

def save_uploaded_file(uploaded_file) -> Optional[str]:
    """Save uploaded file temporarily and return path"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{uploaded_file.name}") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
            
        logger.info(f"Saved {uploaded_file.name} to {tmp_path}")
        return tmp_path
        
    except Exception as e:
        logger.error(f"Failed to save {uploaded_file.name}: {str(e)}")
        return None

def cleanup_temp_file(file_path: str):
    """Clean up temporary file"""
    try:
        if file_path and os.path.exists(file_path):
            os.unlink(file_path)
            logger.info(f"Cleaned up temporary file: {file_path}")
    except Exception as e:
        logger.warning(f"Could not clean up {file_path}: {str(e)}")

def validate_file_content(text: str, filename: str) -> bool:
    """Validate extracted text content"""
    if not text or not text.strip():
        logger.error(f"Empty content from {filename}")
        return False
    
    # Check minimum length
    if len(text.strip()) < 50:
        logger.warning(f"Very short content from {filename} ({len(text)} chars)")
        st.warning(f"⚠️ {filename} contains very little text ({len(text)} characters). Analysis may be limited.")
    
    # Check for common issues
    if text.count('\n') < 3:
        logger.warning(f"Content from {filename} has very few line breaks - may be formatting issue")
    
    return True

def get_file_info(uploaded_file) -> dict:
    """Get file information for display - FIXED size calculation"""
    try:
        # FIXED: Use len(getvalue()) instead of .size attribute
        file_content = uploaded_file.getvalue()
        file_size = len(file_content)
        file_type = uploaded_file.name.split('.')[-1].upper()
        
        if file_size < 1024:
            size_display = f"{file_size} bytes"
        elif file_size < 1024 * 1024:
            size_display = f"{file_size / 1024:.1f} KB"
        else:
            size_display = f"{file_size / (1024 * 1024):.1f} MB"
        
        return {
            'name': uploaded_file.name,
            'size': file_size,
            'size_display': size_display,
            'type': file_type
        }
    except Exception as e:
        logger.error(f"Error getting file info: {str(e)}")
        return {
            'name': uploaded_file.name if uploaded_file else 'Unknown',
            'size': 0,
            'size_display': 'Unknown',
            'type': 'Unknown'
        }
